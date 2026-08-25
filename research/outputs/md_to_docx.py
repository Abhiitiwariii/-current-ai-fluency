#!/usr/bin/env python
"""
md_to_docx.py — convert PRD-full.md into PRD-full.docx.

Handles the deliberately small Markdown subset used by the PRD:
  #/##/### headings, paragraphs, **bold**, [text](url) links,
  - / 1. lists, | pipe tables |, --- rules, and > blockquotes.

Not a general Markdown engine — just enough to render this document cleanly
with real Word heading styles, native tables, and clickable hyperlinks.

Run:  python md_to_docx.py
"""
import re
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "PRD-full.md"
OUT = "PRD-full.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_hyperlink(paragraph, text, url):
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hlink = OxmlElement("w:hyperlink")
    hlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hlink.append(run)
    paragraph._p.append(hlink)


def add_inline(paragraph, text):
    """Render a line with **bold**, *italic*, and [text](url) links."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and not token.startswith("**"):
            paragraph.add_run(token[1:-1]).italic = True
        else:
            m = LINK.fullmatch(token)
            if m:
                add_hyperlink(paragraph, m.group(1), m.group(2))
            else:
                paragraph.add_run(token)


def add_table(doc, rows):
    """rows = list of cell-lists; first is the header."""
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, cell_text in enumerate(row):
            if j >= len(cells):
                continue
            para = cells[j].paragraphs[0]
            add_inline(para, cell_text.strip())
            if i == 0:
                for run in para.runs:
                    run.bold = True


def parse_table_row(line):
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def main():
    lines = open(SRC, encoding="utf-8").read().split("\n")
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule → skip (section spacing handled by headings).
        if stripped == "---":
            i += 1
            continue

        # Table: a pipe line followed by a |---| separator.
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = [parse_table_row(stripped)]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Headings.
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=0)
        # Blockquote.
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run("")
            add_inline(p, stripped[2:])
            for run in p.runs:
                run.italic = True
        # Numbered list.
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        # Bullet list.
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        # Blank line.
        elif stripped == "":
            pass
        # Normal paragraph.
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)

        i += 1

    doc.save(OUT)
    # Report for verification.
    d = docx.Document(OUT)
    print(f"Saved {OUT}: {len(d.paragraphs)} paragraphs, {len(d.tables)} tables")


if __name__ == "__main__":
    main()
